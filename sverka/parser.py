import logging
import re
import traceback
from contextlib import suppress
from datetime import date, datetime
from enum import Enum
from pathlib import Path
from typing import Any, cast, override

import ocrmypdf
import pandas as pd
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from pandas._libs import OutOfBoundsDatetime

from sverka.error import (
    ContractsNofFoundError,
    DataFrameInequalityError,
    DateNotFoundError,
    EmptyTableError,
    ExcesssiveTableCountError,
    InvalidColumnCount,
    JoinPDFNotFoundError,
    JoinProtocolNotFoundError,
    LoanAmountNotFoundError,
    MismatchError,
    ProtocolIDNotFoundError,
    TableNotFound,
    WrongDataInColumnError,
)
from sverka.structures import (
    MONTHS,
    RE_ALPHA_LETTERS,
    RE_COMPLEX_DATE,
    RE_DATE_SEPARATOR,
    RE_END_DATES,
    RE_FILE_CONTENTS,
    RE_FLOAT_NUMBER,
    RE_FLOAT_NUMBER_FULL,
    RE_IBAN,
    RE_JOIN_CONTENTS,
    RE_JOIN_DATE_KAZ,
    RE_JOIN_DATE_RUS,
    RE_JOIN_LOAN_AMOUNT,
    RE_JOIN_PROTOCOL_ID_KAZ,
    RE_JOIN_PROTOCOL_ID_OCR,
    RE_JOIN_PROTOCOL_ID_RUS,
    RE_KZ_LETTERS,
    RE_NUMBER,
    RE_PRIMARY_COLUMN,
    RE_PROTOCOL_ID,
    RE_SECONDARY_COLUMN,
    RE_START_DATE,
    RE_WHITESPACE,
    RE_WRONG_CONTENTS,
)
from sverka.subsidy import Error, ParseJoinContract, ParseSubsidyContract
from utils.db_manager import DatabaseManager
from utils.my_collections import find, index
from utils.office import Office
from utils.utils import compare, get_column_mapping

logger = logging.getLogger("DAMU")


class Backend(Enum):
    PythonDocx = 0
    Docx2Python = 1


def recover_document(file_path: Path) -> Path:
    logger.info(f"Recovering corrupted document: {file_path}")

    og_file_path = file_path.with_name(f"og_{file_path.name}")
    if og_file_path.exists():
        og_file_path.unlink()
    file_path.rename(og_file_path)

    copy_file_path = file_path.parent / f"copy_{file_path.name}"

    try:
        with Office(
            file_path=og_file_path, office_type=Office.Type.WordType
        ) as word:
            word.save_as(copy_file_path, Office.Format.DOCX)
    except (Exception, BaseException) as err:
        og_file_path.unlink()
        copy_file_path.rename(file_path)
        raise err

    og_file_path.unlink()
    copy_file_path.rename(file_path)
    return file_path


def open_document(
    file_path: Path, backend: Backend = Backend.PythonDocx
) -> Any:
    match backend:
        case Backend.PythonDocx:
            try:
                doc = Document(str(file_path))
                return doc
            except Exception:
                logger.warning(
                    f"Failed to open document {file_path}. Attempting recovery..."
                )
                file_path = recover_document(file_path)

                try:
                    return Document(str(file_path))
                except KeyError as err:
                    logger.exception(err)
                    logger.error(
                        f"Failed to open document even after recovery: {err}"
                    )
                    raise err
        case Backend.Docx2Python:
            from docx2python import docx2python

            return docx2python(file_path)


class DamuDocument:
    def __init__(self, file_path: Path) -> None:
        self.file_path = file_path
        self.doc = None
        self.paragraphs: list[str] = []
        self.docx_content = None
        self.is_correct_type = False

    def is_correct_file(self) -> bool | str:
        raise NotImplementedError

    def __repr__(self) -> str:
        return f"{self.__class__.__name__}(file_path={self.file_path.as_posix()}, is_correct_type={self.is_correct_type})"


class SubsidyDocument(DamuDocument):
    def __init__(self, file_path: Path) -> None:
        super().__init__(file_path=file_path)

    @override
    def is_correct_file(self) -> bool | str:
        file_name = self.file_path.name.lower()
        if not file_name.endswith("docx") or file_name.startswith("~$"):
            return False

        self.doc = open_document(self.file_path)

        paragraphs: list[Paragraph] = self.doc.paragraphs

        self.paragraphs = [
            text
            for para in paragraphs
            if (text := RE_WHITESPACE.sub(" ", para.text).strip())
        ]

        if len(self.paragraphs) < 30:
            return False

        fname = file_name.lower()
        if (
            "договор" in fname or "суб" in fname or "дс" in fname
        ) and "присоед" not in fname:
            self.is_correct_type = True
            return self.is_correct_type

        first_n_paras = "\n".join(self.paragraphs[0:10])
        self.is_correct_type = (
            RE_FILE_CONTENTS.search(first_n_paras) is not None
            and RE_WRONG_CONTENTS.search(first_n_paras) is None
        )
        return self.is_correct_type


class JoinDocument(DamuDocument):
    def __init__(self, file_path: Path) -> None:
        super().__init__(file_path=file_path)

    @override
    def is_correct_file(self) -> bool | str:
        file_name = self.file_path.name.lower()
        if not file_name.endswith("docx") or file_name.startswith("~$"):
            return False

        self.doc = open_document(self.file_path)

        self.paragraphs = [
            text
            for para in self.doc.paragraphs
            if (text := RE_WHITESPACE.sub(" ", para.text).strip())
        ]

        fname = file_name.lower()
        if "присоед" in fname:
            self.is_correct_type = True
            return self.is_correct_type

        first_n_paras = "\n".join(self.paragraphs[0:10])
        self.is_correct_type = (
            RE_JOIN_CONTENTS.search(first_n_paras) is not None
        )
        return self.is_correct_type


class TableParser:
    def __init__(self, document: SubsidyDocument | JoinDocument) -> None:
        self.document = document

        self.human_readable = get_column_mapping()
        self.expected_columns = [
            "debt_repayment_date",
            "principal_debt_balance",
            "principal_debt_repayment_amount",
            "agency_fee_amount",
            "recipient_fee_amount",
            "total_accrued_fee_amount",
        ]

    @staticmethod
    def parse_table(
        table: Table, filter_empty: bool = False
    ) -> list[list[str]]:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                text = cell.text.strip().replace("\n", "")
                if filter_empty and not text:
                    continue
                row_data.append(text)
            if any(row_data):
                table_data.append(row_data)
        return table_data

    def find_tables(self) -> list[list[list[str]]]:
        tables: list[list[list[str]]] = []
        for idx, table in enumerate(self.document.doc.tables):
            parsed_table = self.parse_table(table)

            try:
                if len(parsed_table[0]) == 2:
                    continue
            except IndexError:
                pass

            try:
                secondary_column = " ".join(row[1] for row in parsed_table)
            except IndexError:
                continue

            if RE_SECONDARY_COLUMN.search(
                secondary_column
            ) or RE_PRIMARY_COLUMN.search(secondary_column):
                if len(parsed_table) == 1:
                    next_table = self.parse_table(
                        self.document.doc.tables[idx + 1]
                    )
                    parsed_table.extend(next_table)

                tables.append(parsed_table)
                continue

            try:
                next_column = " ".join(row[2] for row in parsed_table)
            except IndexError:
                continue

            if RE_SECONDARY_COLUMN.search(
                next_column
            ) or RE_PRIMARY_COLUMN.search(next_column):
                if len(parsed_table) == 1:
                    next_table = self.parse_table(
                        self.document.doc.tables[idx + 1]
                    )
                    parsed_table.extend(next_table)

                tables.append(parsed_table)
                continue

        return tables

    def get_total_row_idx(self, df: pd.DataFrame) -> int:
        keywords = {"итого", "жиыны", "барлығы", "жиынтығы", "қорытынды"}
        for idx in range(len(df) - 1, -1, -1):
            row = df.iloc[idx]
            for value in row:
                if isinstance(value, str) and any(
                    keyword in value.lower() for keyword in keywords
                ):
                    return idx

        col = df.columns[-1]
        value = cast(str, df.loc[len(df) - 1, col])
        if RE_FLOAT_NUMBER_FULL.search(value) is None:
            for idx in range(len(df) - 2, -1, -1):
                value = cast(str, df.loc[idx, col])
                if RE_FLOAT_NUMBER_FULL.search(value) is not None:
                    return idx

        return len(df) - 1

    def validate_totals(
        self, df: pd.DataFrame, value_columns: list[str]
    ) -> list[str]:
        value_columns.remove("principal_debt_balance")
        is_total_row = df["total"]
        mismatches = []
        for column in value_columns:
            sum_regular_rows = df.loc[~is_total_row, column].sum()
            sum_total_row = df.loc[is_total_row, column].sum()

            if sum_regular_rows != sum_total_row:
                human_readable_column = self.human_readable.get(column)
                message = f"{human_readable_column!r}: {sum_regular_rows} не равно {sum_total_row}"
                mismatches.append(message)

        return mismatches

    def clean_dataframe(self, original_df: pd.DataFrame) -> pd.DataFrame:
        if original_df.empty:
            return original_df

        df: pd.DataFrame = original_df.copy()

        total_row_idx = self.get_total_row_idx(df)
        df = df.iloc[: total_row_idx + 1]

        if len(df.columns) > 6:
            df = df.loc[:, ~df.T.duplicated()]

        df.dropna(axis=1, how="all", inplace=True)

        df = df.loc[:, ~(df == "").all()]

        # if df.loc[0, df.columns[0]].str.isdigit():
        #     # noinspection PyUnresolvedReferences
        #     if (
        #         df.loc[0 : len(df) // 2, df.columns[0]]
        #         .astype(int)
        #         .diff()
        #         .loc[1:]
        #         == 1.0
        #     ).all():
        #         df.drop(df.columns[0], axis=1, inplace=True)

        if len(df.columns) == 6:
            df.columns = self.expected_columns
        else:
            first_row = (
                pd.to_numeric(df[0], errors="coerce").dropna().astype(int)
            )
            if not first_row.empty and first_row.sum() == sum(
                range(1, len(first_row) + 1)
            ):
                df.drop(0, axis=1, inplace=True)
            if not len(df.columns) == 6:
                raise InvalidColumnCount(
                    f"Expected 6 columns - {len(df.columns)} found..."
                )
            else:
                df.columns = self.expected_columns

        df["total"] = False

        if "№" in df.columns:
            df.drop("№", axis=1, inplace=True)

        if df.iloc[0]["debt_repayment_date"] == "1":
            df = df.iloc[1:]
            df.reset_index(inplace=True, drop=True)

        df.loc[total_row_idx, "total"] = True
        df.loc[total_row_idx, "debt_repayment_date"] = None
        df.loc[total_row_idx, "principal_debt_balance"] = None

        df.dropna(axis=1, how="all", inplace=True)

        try:
            df.loc[:, "debt_repayment_date"] = pd.to_datetime(
                df.loc[:, "debt_repayment_date"], dayfirst=True, format="mixed"
            )
        except (OutOfBoundsDatetime, ValueError) as err:
            raise err

        df["debt_repayment_date"] = df["debt_repayment_date"].astype(
            "datetime64[ns]"
        )

        try:
            if not all(
                df.loc[0 : total_row_idx - 1, "principal_debt_balance"]
                .astype("datetime64[ns]")
                .isna()
            ):
                raise WrongDataInColumnError(
                    "Не удалось перевести суммы в числовые значения. Даты в колонке под названием 'Сумма остатка основного долга'."
                )
        except Exception:
            pass

        columns_to_process = [
            "principal_debt_balance",
            "principal_debt_repayment_amount",
            "agency_fee_amount",
            "recipient_fee_amount",
            "total_accrued_fee_amount",
        ]

        df[columns_to_process] = (
            (
                df[columns_to_process]
                .replace({"-": "0", "": "0", "[  ]+": "", ",": "."}, regex=True)
                .astype(float)
                * 100
            )
            .round()
            .astype("Int64")
        )

        # df = df.where(pd.notna(df), None)
        df.reset_index(inplace=True, drop=True)

        if mismatches := self.validate_totals(
            df, value_columns=columns_to_process
        ):
            raise MismatchError("\n".join(mismatches))

        return df

    def parse_tables(
        self, contract: ParseSubsidyContract | ParseJoinContract
    ) -> list[pd.DataFrame]:
        tables = self.find_tables()
        table_count = len(tables)

        if not tables:
            raise TableNotFound(
                self.document.file_path.name,
                contract.contract_id,
                target="График погашения",
            )

        if table_count < 1 or table_count > 2:
            raise ExcesssiveTableCountError(
                self.document.file_path.name, contract.contract_id, table_count
            )

        logger.info(f"PARSE - found {table_count} table")

        dfs = []
        for table in tables:
            data_start_row_idx = index(
                items=table,
                condition=lambda row: not any(
                    RE_ALPHA_LETTERS.search(cell) is not None for cell in row
                ),
                default=None,
            )

            if not data_start_row_idx:
                raise EmptyTableError()

            data_start_row_idx2 = index(
                items=table,
                condition=lambda row: any(
                    len(cell) > 1 and not RE_ALPHA_LETTERS.search(cell)
                    for cell in row
                ),
            )

            if data_start_row_idx != data_start_row_idx2:
                logger.info(f"{data_start_row_idx=}, {data_start_row_idx2=}")
                data_start_row_idx = data_start_row_idx2

            df = pd.DataFrame(table[data_start_row_idx:])
            df = self.clean_dataframe(df)
            dfs.append(df)
        return dfs


class Parser:
    def __init__(
        self,
        document: SubsidyDocument | JoinDocument,
        contract: ParseSubsidyContract | ParseJoinContract,
    ) -> None:
        self.contract = contract
        self.document = document

        self.table_parser = TableParser(document=self.document)

    def find_ibans(self) -> list[str]:
        ibans: list[str] = RE_IBAN.findall(
            "".join(
                self.document.paragraphs[
                    1 : int(len(self.document.paragraphs) * 0.7)
                ]
            )
        )
        if ibans:
            ibans = [iban.replace('"', "") for iban in ibans]
            return ibans

        if not self.document.docx_content:
            self.document.docx_content = open_document(
                file_path=self.document.file_path, backend=Backend.Docx2Python
            )

        ibans = RE_IBAN.findall(self.document.docx_content.text)
        return ibans

    def find_dbz(self) -> tuple[str | None, date | None]:
        dbz_data = None
        for table in self.document.doc.tables:
            parsed_table = self.table_parser.parse_table(table)
            if not parsed_table:
                continue
            first_row = parsed_table[0]
            if not first_row:
                continue
            first_col = first_row[0]
            if RE_KZ_LETTERS.search(first_col):
                continue

            if (
                first_col.count("/") == 2
                or "Договор банковского займа" in first_col
            ):
                value = first_row[1]
                dbz_data = value.strip()
                break

        if not dbz_data:
            return None, None

        easy_exprs = [
            r"Заявление[№ ]+([^ ]+) на выдачу банковского займа от (\d+.\d+.\d+)",
            r"ДОГОВОР БАНКОВСКОГО ЗАЙМА № ?([^ ]+) от (\d+.\d+.\d+)",
            r"к заявлению на выдачу банковского займа[№ ]+([^ ]+) +от (\d+.\d+.\d+)",
            r"займа[№ ]+([^ ]+) от (\d+.\d+.\d+)",
            r"^[№ ]+([^ ]+)\s*от\s+(\d+.\d+.\d+)",
            r"Заявление[№ ]+([^ ]+)\s+на\s+выдачу[А-Яа-я ]+(\d+.\d+.\d+)",
            r"Соглашение\s+об\s+открытии\s+кредитной\s+линии[№ ]+([^ ]+)\sот\s(\d+.\d+.\d+)",
            r"Заявление\s[№ ]+([^ ]+)\s*от\s*(\d+.\d+.\d+)"
            r"ЗАЯ\w+\s+О\s+ПРИС\w+\s+[№ ]+([^ ]+)\s+к\s+Дого\w+\s+прис\w+\s+\(о\s+пре\w+\s+бан\w+\s+Зай\w+\s+вне\s+КЛ/ЛК\)\s+от\s+(\d+.\d+.\d+)",
            r"Акцес\w+\s+Дог\w+[№ ]+([^ ]+)\s+\(о\s+пред\w+\s+банк\w+\s+за\w+\)\s+от\s+(\d+.\d+.\d+)",
            r"График\s+погашения\s+кредита\s+от\s+(\d+.\d+.\d+)\s+г\.\s+к\s+заявлению\s+[№ ]+([^ ]+)",
            r"Договор\s+лизинга[№ ]+([^ ]+)\s+от\s+(\d+.\d+.\d+)",
            r"Договор\s+финансового\s+лизинга[№ ]+([^ ]+)\s+от\s+(\d+.\d+.\d+)",
            r"Заявление\s+[№ ]+([^ ]+)\s*от\s*(\d+.\d+.\d+)",
        ]

        hard_exprs = [
            r"ДОГОВОР БАНКОВСКОГО ЗАЙМА № ?([^ ]+) \(.+\) от [«\"]?(\d+)[»\"]? (\w+) (\d+)",
            r"Заявление о присоединении[№ ]+([^ ]+) от [«\"]?(\d+)[»\"]? (\w+) (\d+)",
            r"ДОГОВОР БАНКОВСКОГО ЗАЙМА[№ ]+([^ ]+) в рамках соглашения кл/лк[ \w]*[«\"]?(\d+)[»\"]? (\w+) (\d+)",
            r"^[№ ]+([^ ]+)\s*от\s+[«\"]?(\d+)[»\"]?\s+(\w+)\s+(\d+)",
            r"договор\w? банков\w+ займа[№ ]+([^ ]+)\s+от\s+[«\"]?(\d+)[»\"]? (\w+) (\d+)",
            r"Соглашение\s+о\s+предоставлении\s+кредитной\s+линии[№ ]+([^ ]+)\s+от\s+[«\"]?(\d+)[»\"]? (\w+) (\d+)",
            r"Соглашение об открытии кредитной линии[№ ]+([^ ]+)\sот\s[«\"]?(\d+)[»\"]? (\w+) (\d+)",
            r"ЗАЯВЛ\w+\s+О\s+ПРИСОЕ\w+[№ ]+([^ ]+)\s+к догов\w+ присое\w+.+?от [«\"]?(\d+)[»\"]? (\w+) (\d+)",
            r"Акцес\w+\s+Дог\w+[№ ]+([^ ]+)\s*от\s*[«\"]?(\d+)[»\"]? (\w+) (\d+)",
        ]

        easy_pats = [re.compile(expr, re.IGNORECASE) for expr in easy_exprs]
        hard_pats = [re.compile(expr, re.IGNORECASE) for expr in hard_exprs]

        match = next(
            (m for pat in easy_pats if (m := pat.search(dbz_data))), None
        )
        if match:
            dbz_id, dbz_date_str = match.groups()

            if "." in dbz_id:
                dbz_id, dbz_date_str = dbz_date_str, dbz_id

            fmt = "%d.%m.%y" if len(dbz_date_str) == 8 else "%d.%m.%Y"
            try:
                dbz_date = datetime.strptime(dbz_date_str, fmt).date()
            except ValueError:
                dbz_date = datetime.strptime(dbz_date_str, "%d/%m/%Y").date()
        else:
            match = next(
                (m for pat in hard_pats if (m := pat.search(dbz_data))), None
            )
            if match:
                dbz_id, day, month, year = match.groups()
                month = MONTHS.get(month[0:3])
                fmt = "%d.%m.%y" if len(year) == 2 else "%d.%m.%Y"
                dbz_date = datetime.strptime(
                    f"{day}.{month}.{year}", fmt
                ).date()
            else:
                return dbz_data, None

        return dbz_id, dbz_date

    def find_subsidy_loan_amount(self) -> float | None:
        for table in self.document.doc.tables:
            parsed_table = self.table_parser.parse_table(
                table, filter_empty=True
            )
            row_count = len(parsed_table)
            if not (
                7 <= row_count <= 9
                and sum(1 for row in parsed_table if len(row) == 2)
                > row_count // 2
            ):
                continue

            row_idx = index(
                parsed_table,
                condition=lambda row: row[0].startswith("Сумма"),
                default=3,
            )

            amount_str = parsed_table[row_idx][1]
            match = RE_FLOAT_NUMBER.search(amount_str)
            if not match:
                continue

            match_str = match.group(1)
            match_str = (
                match_str.replace(" ", "").replace(",", ".").replace(" ", "")
            )
            amount = float(match_str)
            return amount


class SubsidyParser(Parser):
    def __init__(
        self, document: SubsidyDocument, contract: ParseSubsidyContract
    ) -> None:
        super().__init__(document=document, contract=contract)

    def find_subsidy_date(self, pat: re.Pattern) -> date | None:
        para = find(
            self.document.paragraphs,
            condition=lambda p: pat.search(p) is not None,
        )

        if not para or (isinstance(para, str) and len(para) < 30):
            if not self.document.docx_content:
                self.document.docx_content = open_document(
                    file_path=self.document.file_path,
                    backend=Backend.Docx2Python,
                )

            new_pat = re.compile(pat.pattern.replace(r"\.", "[.)]"))
            para = ""
            for line in self.document.docx_content.text.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if new_pat.search(line):
                    para += " " + line

        if "ислам" in para:
            raise ValueError("Договор Исламского банка...")

        para = (
            para.replace('"', " ")
            .replace("«", " ")
            .replace("»", " ")
            .replace(" .", ".")
            .replace("-ін", " ін")
            .replace("г.", " г.")
            .replace("ж.", " ж.")
            .replace("года", " года")
            .replace("жыл", " жыл")
        )

        for month in MONTHS.keys():
            para = para.replace(month, f" {month}")

        date_str = (
            match.group(1) if (match := RE_COMPLEX_DATE.search(para)) else None
        )

        if not isinstance(date_str, str):
            raise DateNotFoundError(
                self.document.file_path.name, self.contract.contract_id, para
            )

        date_str = date_str.replace("-", ".").replace("/", ".")

        with suppress(ValueError):
            if len(date_str) == 10:
                res = datetime.strptime(date_str, "%d.%m.%Y").date()
                return res
            elif len(date_str) == 8:
                res = datetime.strptime(date_str, "%d.%m.%y").date()
                return res

        items: tuple[str, ...] = tuple(
            item
            for item in RE_DATE_SEPARATOR.split(date_str)
            if item
            and all(not item.startswith(word) for word in {"год", "жыл"})
            and (item.isdigit() or len(item) > 1)
        )

        if len(items) != 3:
            raise DateNotFoundError(
                self.document.file_path.name, self.contract.contract_id, para
            )

        if len(items[0]) == 2:
            day, month, year = items
        else:
            year, day, month = items

        if not day.isdigit():
            day, month = month, day

        if not month.isdigit():
            month = month[0:3]
            month_num = MONTHS.get(month)
        else:
            month_num = month

        if month_num is None:
            month_num = month

        if not year.isdigit():
            year_match = RE_NUMBER.search(year)
            if year_match:
                year = year_match.group(1)

        fmt = "%d.%m.%Y" if len(year) == 4 else "%d.%m.%y"

        try:
            res = datetime.strptime(f"{day}.{month_num}.{year}", fmt).date()
            return res
        except ValueError as err:
            raise err

    def find_subsidy_protocol_id(self) -> str | None:
        termin_para_idx = index(
            self.document.paragraphs,
            condition=lambda p: "ермин" in p,
            default=-1,
        )

        if termin_para_idx != -1:
            text = [
                x
                for x in "".join(
                    self.document.paragraphs[:termin_para_idx]
                ).split(";")
                if x
            ][-1]

            protocol_ids = RE_PROTOCOL_ID.findall(text)
            if not protocol_ids:
                raise ProtocolIDNotFoundError("Протокол не найден")
            return protocol_ids[-1]

        if not self.document.docx_content:
            self.document.docx_content = open_document(
                file_path=self.document.file_path, backend=Backend.Docx2Python
            )

        full_text = self.document.docx_content.text
        termin_idx = full_text.find("ермин")
        protocol_ids = RE_PROTOCOL_ID.findall(full_text[:termin_idx])
        if not protocol_ids:
            raise ProtocolIDNotFoundError("Протокол не найден")
        return protocol_ids[-1]

    def parse_document(self) -> list[pd.DataFrame]:
        self.contract.file_name = self.document.file_path.name

        self.contract.protocol_id = self.find_subsidy_protocol_id()
        if not self.contract.protocol_id:
            logger.error("PARSE - WARNING - protocols not found")
            raise ValueError("Protocol IDs not found...")

        ibans = self.find_ibans()
        if len(set(ibans)) > 1:
            logger.error(
                f"PARSE - Different IBAN codes found in the document: {ibans}"
            )
            raise ValueError(f"IBANs are different - {ibans!r}")
        if ibans:
            self.contract.iban = ibans[0]

        self.contract.dbz_id, self.contract.dbz_date = self.find_dbz()

        if not self.contract.iban:
            logger.error("PARSE - WARNING - IBAN not found")
            raise ValueError("IBANs not found...")
        else:
            logger.info(f"PARSE - iban={self.contract.iban!r}")

        self.contract.start_date = self.find_subsidy_date(RE_START_DATE)
        logger.info(f"PARSE - start_date={self.contract.start_date}")
        if not self.contract.start_date:
            raise DateNotFoundError(
                self.document.file_path.name, self.contract.contract_id
            )

        last_err = None
        for pat in RE_END_DATES:
            try:
                self.contract.end_date = self.find_subsidy_date(pat)
                break
            except (Exception, BaseException) as err:
                last_err = err

        if last_err and not self.contract.end_date:
            raise last_err

        if not self.contract.end_date:
            raise DateNotFoundError(
                self.document.file_path.name, self.contract.contract_id
            )

        logger.info(f"PARSE - end_date={self.contract.end_date}")

        self.contract.loan_amount = self.find_subsidy_loan_amount()
        if not self.contract.loan_amount:
            logger.error("PARSE - WARNING - loan_amount=None")
        else:
            logger.info(f"PARSE - loan_amount={self.contract.loan_amount!r}")

        dfs = self.table_parser.parse_tables(self.contract)

        return dfs


class JoinParser(Parser):
    def __init__(
        self, document: JoinDocument, contract: ParseJoinContract
    ) -> None:
        super().__init__(document=document, contract=contract)

    def find_join_protocol_id_loan_amount(self) -> tuple[str, float]:
        from pypdf import PdfReader

        document_folder = self.document.file_path.parent
        pdf_path = next(
            (
                f
                for f in document_folder.iterdir()
                if f.name.endswith("pdf") and "заявлен" in f.name.lower()
            ),
            None,
        )
        if not pdf_path:
            raise JoinPDFNotFoundError(
                "PDF файл 'Заявление получателя к договору присоединения' не найден"
            )

        protocol_id = None
        loan_amount = None
        reader = PdfReader(pdf_path)

        pat_rus, pat_kaz = (
            RE_JOIN_PROTOCOL_ID_RUS,
            RE_JOIN_PROTOCOL_ID_KAZ,
        )
        pat_loan_amount = RE_JOIN_LOAN_AMOUNT

        for page in reader.pages:
            if protocol_id and loan_amount:
                break

            text = page.extract_text().replace("\n", " ")
            text = re.sub(r"\s+", " ", text)
            if not text:
                continue

            if (
                match := (pat_rus.search(text) or pat_kaz.search(text))
            ) is not None:
                protocol_id = match.group(1)

            search1, search2 = (
                text.find("умма кредита"),
                text.find("редит/лизинг сомасы"),
            )
            if search1 != -1 or search2 != -1:
                if search1 != -1 and search2 == -1:
                    search_idx = search1
                elif search2 != -1 and search1 == -1:
                    search_idx = search2
                else:
                    continue

                snippet = text[search_idx : search_idx + 100]
                if (match := pat_loan_amount.search(snippet)) is not None:
                    match_str = match.group(1).strip()
                    if "," in match_str and "." in match_str:
                        match_str = match_str.rsplit(" ", maxsplit=1)[0]

                    loan_amount = float(
                        match_str.replace(" ", "").replace(",", ".")
                    )

        if not protocol_id:
            temp_folder = Path("C:/Temp") / self.contract.contract_id
            temp_folder.mkdir(exist_ok=True, parents=True)

            output_txt_path = temp_folder / "output.txt"
            if not output_txt_path.exists():
                ocrmypdf.ocr(
                    pdf_path,
                    (temp_folder / "output.pdf"),
                    sidecar=output_txt_path,
                    language="eng+rus+kaz",
                    deskew=True,
                    force_ocr=True,
                )

            with output_txt_path.open(encoding="utf-8") as f:
                for line in f.readlines():
                    if protocol_id and loan_amount:
                        break

                    line = line.strip().lower()
                    if (
                        "номер и дата решения" in line
                        and (match := RE_JOIN_PROTOCOL_ID_OCR.search(line))
                        is not None
                    ):
                        protocol_id = match.group(1)

                    if loan_amount:
                        continue

                    search1, search2 = -1, -1
                    if (search1 := line.find("умма кредита")) or (
                        search2 := line.find("редит/лизинг сомасы")
                    ):
                        if search1 != -1 and search2 == -1:
                            search_idx = search1
                        elif search2 != -1 and search1 == -1:
                            search_idx = search2
                        else:
                            continue
                        snippet = line[search_idx : search_idx + 100]
                        match = RE_JOIN_LOAN_AMOUNT.search(snippet)
                        if match:
                            loan_amount = float(
                                match.group(1)
                                .strip()
                                .replace(" ", "")
                                .replace(",", ".")
                            )

            if not protocol_id:
                raise JoinProtocolNotFoundError(
                    f"Номер протокола не найден в файле {pdf_path.name!r}"
                )

            if not loan_amount:
                raise LoanAmountNotFoundError(
                    f"Сумма кредита не найдена в файле {pdf_path.name}"
                )

        return protocol_id, loan_amount

    def find_join_dates(self) -> tuple[date, date]:
        pat_rus, pat_kaz = (
            RE_JOIN_DATE_RUS,
            RE_JOIN_DATE_KAZ,
        )

        dates: list[date] = []
        for para in self.document.paragraphs:
            if len(dates) == 2:
                break

            if (
                match := (pat_rus.search(para) or pat_kaz.search(para))
            ) is not None:
                date_str = match.group(1)
                date_obj = datetime.strptime(date_str, "%d.%m.%Y").date()
                dates.append(date_obj)

        if dates[1] > dates[0]:
            return dates[0], dates[1]
        else:
            return dates[1], dates[0]

    def parse_document(self) -> list[pd.DataFrame]:
        self.contract.file_name = self.document.file_path.name

        self.contract.protocol_id, self.contract.loan_amount = (
            self.find_join_protocol_id_loan_amount()
        )

        if not self.contract.protocol_id:
            logger.error("PARSE - WARNING - protocol_id=None")
        else:
            logger.info(f"PARSE - protocol_id={self.contract.protocol_id!r}")

        if not self.contract.loan_amount:
            logger.error("PARSE - WARNING - loan_amount=None")
        else:
            logger.info(f"PARSE - loan_amount={self.contract.loan_amount!r}")

        ibans = self.find_ibans()
        if len(set(ibans)) > 1:
            logger.error(
                f"PARSE - Different IBAN codes found in the document: {ibans}"
            )
            raise ValueError(f"IBANs are different - {ibans!r}")
        if ibans:
            self.contract.iban = ibans[0]

        self.contract.dbz_id, self.contract.dbz_date = self.find_dbz()

        if not self.contract.iban:
            logger.error("PARSE - WARNING - IBAN not found")
            raise ValueError("IBANs not found...")
        else:
            logger.info(f"PARSE - iban={self.contract.iban!r}")

        self.contract.start_date, self.contract.end_date = (
            self.find_join_dates()
        )
        logger.info(f"PARSE - start_date={self.contract.start_date}")
        logger.info(f"PARSE - end_date={self.contract.start_date}")
        if not self.contract.start_date or not self.contract.end_date:
            raise DateNotFoundError(
                self.document.file_path.name, self.contract.contract_id
            )

        dfs = self.table_parser.parse_tables(self.contract)

        return dfs


def parse_document(
    contract_id: str,
    contract_type: str,
    download_folder: Path,
    db: DatabaseManager,
) -> ParseSubsidyContract:
    contract = ParseSubsidyContract(
        contract_id=contract_id, error=Error(contract_id=contract_id)
    )

    documents_folder = download_folder / contract_id / "documents"

    documents: list[SubsidyDocument | JoinDocument]
    if contract_type in [
        "Первый график к договору присоединения",
        "Транш к договору присоединения",
    ]:
        document_cls = JoinDocument
        parser_cls = JoinParser
    else:
        document_cls = SubsidyDocument
        parser_cls = SubsidyParser

    try:
        documents = [
            doc
            for fpath in documents_folder.iterdir()
            if (doc := document_cls(fpath)).is_correct_file()
        ]
    except (KeyError, ValueError, FileNotFoundError) as err:
        logger.exception(err)
        if contract.error:
            contract.error.traceback = f"{err!r}\n{traceback.format_exc()}"
            contract.error.human_readable = contract.error.get_human_readable()
            contract.error.save(db)
        contract.save(db)
        return contract

    document_count = len(documents)

    if documents:
        dfs: list[pd.DataFrame] = []
        for jdx, document in enumerate(documents, start=1):
            parser = parser_cls(document=document, contract=contract)

            if document_count > 1:
                logger.info(
                    f"PARSE - {jdx}/{len(documents)} - {document.file_path.name!r}"
                )
            try:
                dfs.extend(parser.parse_document())
            except (Exception, BaseException) as err:
                logger.exception(err)
                logger.error(f"{err!r}")
                if contract.error:
                    contract.error.traceback = (
                        f"{err!r}\n{traceback.format_exc()}"
                    )
                    contract.error.error = err
                    contract.error.human_readable = (
                        contract.error.get_human_readable()
                    )
                    contract.error.save(db)
                contract.save(db)
                return contract
        try:
            if len(dfs) == 2 and not compare(dfs[0], dfs[1]):
                raise DataFrameInequalityError(
                    "DataFrames not equal to each other"
                )
        except DataFrameInequalityError as err:
            logger.exception(err)
            logger.error(f"{err!r}")
            if contract.error:
                contract.error.traceback = f"{err!r}\n{traceback.format_exc()}"
                contract.error.human_readable = (
                    contract.error.get_human_readable()
                )
                contract.error.save(db)
            contract.save(db)
            return contract

        if len(dfs):
            contract.df = dfs[0]
    else:
        try:
            raise ContractsNofFoundError(
                f"EDO - {contract_id} does not have subsidy contracts..."
            )
        except ContractsNofFoundError as err:
            logger.exception(err)
            logger.error(f"{err!r}")
            contract.error.traceback = f"{err!r}\n{traceback.format_exc()}"
            contract.error.human_readable = contract.error.get_human_readable()

    if contract.df is not None:
        contract.settlement_date = int(
            contract.df["debt_repayment_date"].dt.day.value_counts().idxmax()
        )

    contract.save(db)
    contract.error.save(db)

    return contract
