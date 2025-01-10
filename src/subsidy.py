import logging
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Optional

from docx import Document
from docx.document import Document as DocumentObject

from src.edo import ProtocolIDs, SubsidyContract
from src.utils.custom_list import CustomList
from src.utils.office import Word, WordDocument


@dataclass
class RegexPatterns:
    file_name: re.Pattern = re.compile(
        r"((дог\w*.?.суб\w*.?)|(дс))",
        re.IGNORECASE,
    )
    file_contents: re.Pattern = re.compile(
        r"((бір бөлігін субсидиялау туралы)|(договор субсидирования))",
        re.IGNORECASE,
    )
    protocol_id: re.Pattern = re.compile(r"№.?(\d{6})")


def open_document(file_path: Path) -> DocumentObject:
    try:
        doc = Document(str(file_path))
    except KeyError:
        logging.warning(f"Corrupted document. Attempting to re-save it...")

        _file_path = file_path

        og_file_path = file_path.with_name(f"og_{_file_path.name}")
        file_path.rename(og_file_path)

        copy_file_path = file_path.parent / f"copy_{_file_path.name}"
        with Word() as word:
            with WordDocument(word, og_file_path) as word_doc:
                word_doc.save_as(copy_file_path, 16)

        og_file_path.unlink()
        copy_file_path.rename(_file_path)
        file_path = _file_path
        doc = Document(str(file_path))

    return doc


class Paragraphs(CustomList[str]):
    pass


class SubsidyParser:
    def __init__(
        self, contract: SubsidyContract, regex_patterns: RegexPatterns
    ) -> None:
        self.contract_folder = Path(contract.save_location).parent
        self.regex_patterns = regex_patterns
        self.file_path: Optional[Path] = None
        self._doc: Optional[DocumentObject] = None
        self._paragraphs: Paragraphs = Paragraphs()

    def find_subsidy_contact_file(self) -> Optional[Path]:
        for file_path in self.contract_folder.iterdir():
            self.file_path = file_path
            if file_path.name.endswith("docx") and self.is_subsidy_contract_file(
                file_path
            ):
                return self.file_path
        self.file_path = None
        return None

    def is_subsidy_contract_file(self, file_path: Path) -> bool:
        if self.regex_patterns.file_name.search(file_path.name):
            return True

        if self.regex_patterns.file_contents.search("\n".join(self.paragraphs[0:10])):
            return True

        self._paragraphs.clear()
        self._doc = None
        return False

    @property
    def doc(self) -> DocumentObject:
        if not self._doc:
            self._doc = open_document(self.file_path)
            self._doc = Document(str(self.file_path))
        return self._doc

    @property
    def paragraphs(self) -> Paragraphs:
        if not self._paragraphs:
            self._paragraphs = Paragraphs(
                [
                    text
                    for para in self.doc.paragraphs
                    if (text := re.sub(r"\s+", " ", para.text.lower()).strip())
                ]
            )
        return self._paragraphs

    def find_protocol_ids(self) -> ProtocolIDs:
        protocol_ids = ProtocolIDs()
        termin_para_idx = self.paragraphs.index(condition=lambda p: "термин" in p)

        if not termin_para_idx:
            logging.error(f"EDO - no protocol ids found...")
            return protocol_ids

        para = self.paragraphs[termin_para_idx - 1].split(";")[-1]

        protocol_ids.items = self.regex_patterns.protocol_id.findall(para)

        return protocol_ids
